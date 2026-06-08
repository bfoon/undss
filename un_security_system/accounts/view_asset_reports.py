# view_asset_reports.py
"""
Comprehensive Asset Management Reports
──────────────────────────────────────
Accessible at accounts:asset_reports  (ICT / Ops / Superuser only)

Covers:
  • Asset inventory & status breakdown
  • Asset requests pipeline
  • Assets issued (assignments) with holder details
  • Mobile lines & line allocations
  • Consumable supplies & supply requests
  • Allocation summaries by unit / category
  • Historical trends (monthly)

Downloads:
  • /asset-reports/download/excel/  → openpyxl workbook (multi-sheet)
  • /asset-reports/download/word/   → python-docx Word document
"""

import io
import json
from datetime import date, timedelta
from collections import defaultdict

from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db.models import Count, F, Q, Sum
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.utils import timezone

from .models import (
    AgencyAssetRoles, AgencyServiceConfig, Asset, AssetCategory,
    AssetHistory, AssetRequest, AssetReturnRequest,
    ConsumableItem, ConsumableRequest, ConsumableRequestItem,
    ConsumableStockLog, MobileLine, MobileLineReactivationRequest, Unit, User,
)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _require_manager(user, agency):
    """
    Return True if the user may view reports.
    Allowed: superuser, ict_focal role, ICT custodian (via AgencyAssetRoles),
             Operations Manager (via AgencyAssetRoles), unit heads / asset managers.
    """
    if not user.is_authenticated:
        return False

    # Superuser always allowed
    if user.is_superuser:
        return True

    # Role-based shortcut (ict_focal flag on the user)
    if getattr(user, "role", "") == "ict_focal":
        return True

    # Fetch roles directly from DB — never rely on the cached reverse relation
    try:
        roles = AgencyAssetRoles.objects.get(agency=agency)
    except AgencyAssetRoles.DoesNotExist:
        roles = None

    if roles:
        # ICT custodian M2M
        if roles.ict_custodian.filter(id=user.id).exists():
            return True
        # Operations Manager FK
        if roles.operations_manager_id and roles.operations_manager_id == user.id:
            return True

    # Unit head or asset manager for any unit in this agency
    from .models import Unit as _Unit
    if _Unit.objects.filter(agency=agency).filter(
        Q(unit_head=user) | Q(asset_managers=user)
    ).exists():
        return True

    return False


def _months_range(n=12):
    """Return list of (year, month, label) for the last n months."""
    today = date.today()
    months = []
    for i in range(n - 1, -1, -1):
        d = (today.replace(day=1) - timedelta(days=i * 28)).replace(day=1)
        months.append((d.year, d.month, d.strftime("%b %Y")))
    return months


def _build_report_data(agency):
    """Build all report datasets. Returns a big dict consumed by template + downloads."""
    now = timezone.now()
    today = now.date()

    # ── Assets ──────────────────────────────────────────────────────────────
    assets_qs = Asset.objects.filter(agency=agency).select_related(
        "category", "unit", "current_holder"
    )

    asset_status_counts = dict(
        assets_qs.values("status").annotate(c=Count("id")).values_list("status", "c")
    )
    total_assets = assets_qs.count()
    available_assets = asset_status_counts.get("available", 0)
    assigned_assets = asset_status_counts.get("assigned", 0)
    retired_assets = asset_status_counts.get("retired", 0)
    maintenance_assets = asset_status_counts.get("maintenance", 0)

    # By category
    assets_by_cat = list(
        assets_qs.values("category__name").annotate(
            total=Count("id"),
            available=Count("id", filter=Q(status="available")),
            assigned=Count("id", filter=Q(status="assigned")),
            retired=Count("id", filter=Q(status="retired")),
        ).order_by("-total")
    )

    # By unit
    assets_by_unit = list(
        assets_qs.values("unit__name").annotate(
            total=Count("id"),
            available=Count("id", filter=Q(status="available")),
            assigned=Count("id", filter=Q(status="assigned")),
        ).order_by("-total")
    )
    for row in assets_by_unit:
        if not row["unit__name"]:
            row["unit__name"] = "Unallocated / Core"

    # EOL assets
    eol_assets = [a for a in assets_qs if getattr(a, "is_eol_due", False) and a.status != "retired"]

    # ── Asset Requests ───────────────────────────────────────────────────────
    requests_qs = AssetRequest.objects.filter(agency=agency).select_related(
        "requester", "category", "unit", "assigned_asset"
    )

    req_status_counts = dict(
        requests_qs.values("status").annotate(c=Count("id")).values_list("status", "c")
    )
    total_requests = requests_qs.count()

    # By category
    # AssetRequest statuses: draft, pending_manager, pending_ict, assigned, received, rejected, cancelled
    requests_by_cat = list(
        requests_qs.values("category__name").annotate(
            total=Count("id"),
            fulfilled=Count("id", filter=Q(status="received")),
            pending=Count("id", filter=Q(status__in=["pending_manager", "pending_ict", "assigned"])),
            cancelled=Count("id", filter=Q(status="cancelled")),
            rejected=Count("id", filter=Q(status="rejected")),
        ).order_by("-total")
    )

    # Monthly request trend (last 12 months)
    months = _months_range(12)
    monthly_requests = []
    for y, m, label in months:
        c = requests_qs.filter(created_at__year=y, created_at__month=m).count()
        monthly_requests.append({"label": label, "count": c})

    # Avg fulfillment time (created -> received/verified)
    # AssetRequest has no fulfilled_at; "received" status + requester_verified_at is the terminal state
    received_qs = requests_qs.filter(status="received", requester_verified_at__isnull=False)
    if received_qs.exists():
        deltas = [
            (r.requester_verified_at - r.created_at).days
            for r in received_qs
            if r.requester_verified_at and r.created_at
        ]
        avg_fulfillment_days = round(sum(deltas) / len(deltas), 1) if deltas else None
    else:
        avg_fulfillment_days = None

    # ── Assets Issued (assignments) ──────────────────────────────────────────
    issued_assets = list(
        assets_qs.filter(status="assigned").select_related(
            "current_holder", "category", "unit"
        ).order_by("unit__name", "category__name", "current_holder__last_name")
    )

    # Pending returns
    pending_returns = list(
        AssetReturnRequest.objects.filter(agency=agency, status="pending_ict").select_related(
            "asset", "asset__category", "requested_by"
        ).order_by("-created_at")
    )

    # ── Mobile Lines ─────────────────────────────────────────────────────────
    lines_qs = MobileLine.objects.filter(agency=agency).select_related(
        "assigned_to", "custodian"
    )

    line_status_counts = dict(
        lines_qs.values("status").annotate(c=Count("id")).values_list("status", "c")
    )
    total_lines = lines_qs.count()

    # By type
    lines_by_type = list(
        lines_qs.values("line_type").annotate(
            total=Count("id"),
            assigned=Count("id", filter=Q(status="assigned")),
            available=Count("id", filter=Q(status="available")),
            suspended=Count("id", filter=Q(status="suspended")),
        ).order_by("-total")
    )

    # Allocation by unit (from assigned_to's unit)
    line_by_unit = list(
        lines_qs.filter(status="assigned").values(
            "assigned_to__unit__name"
        ).annotate(count=Count("id")).order_by("-count")
    )
    for row in line_by_unit:
        if not row["assigned_to__unit__name"]:
            row["assigned_to__unit__name"] = "Unallocated"

    # Lines details
    all_lines = list(lines_qs.order_by("status", "line_type", "msisdn"))

    # Reactivation requests
    reactivation_requests = list(
        MobileLineReactivationRequest.objects.filter(agency=agency).select_related(
            "line", "requested_by"
        ).order_by("-created_at")[:30]
    )

    # ── Consumables / Supplies ───────────────────────────────────────────────
    consumables_qs = ConsumableItem.objects.filter(agency=agency, is_active=True).select_related("category")

    stock_summary = list(
        consumables_qs.values("category__name").annotate(
            total_items=Count("id"),
            total_stock=Sum("stock_qty"),
            low_stock_items=Count("id", filter=Q(stock_qty__lte=F("low_stock_threshold"))),
        ).order_by("category__name")
    )

    all_consumables = list(consumables_qs.order_by("category__name", "name"))

    creqs_qs = ConsumableRequest.objects.filter(agency=agency).select_related(
        "requester", "unit"
    ).prefetch_related("line_items__item")

    creq_status_counts = dict(
        creqs_qs.values("status").annotate(c=Count("id")).values_list("status", "c")
    )
    total_creqs = creqs_qs.count()

    # Top requested consumables (from requests that reached dispatch/fulfillment stage)
    # ConsumableRequest statuses: pending, approved, partially_fulfilled, fulfilled, rejected, cancelled
    top_consumables = list(
        ConsumableRequestItem.objects.filter(
            request__agency=agency,
            request__status__in=["approved", "partially_fulfilled", "fulfilled"]
        ).values("item__name", "item__category__name").annotate(
            total_qty=Sum("quantity_requested"),
            request_count=Count("id"),
        ).order_by("-total_qty")[:15]
    )

    # Monthly supply requests trend
    monthly_creqs = []
    for y, m, label in months:
        c = creqs_qs.filter(created_at__year=y, created_at__month=m).count()
        monthly_creqs.append({"label": label, "count": c})

    # Stock logs (recent 30)
    recent_stock_logs = list(
        ConsumableStockLog.objects.filter(agency=agency).select_related(
            "item", "actor"
        ).order_by("-created_at")[:30]
    )

    # ── Allocation Summary ───────────────────────────────────────────────────
    units_qs = Unit.objects.filter(agency=agency).select_related("unit_head")

    allocation_by_unit = []
    for unit in units_qs.order_by("name"):
        unit_assets = assets_qs.filter(unit=unit)
        unit_lines = lines_qs.filter(assigned_to__unit=unit)
        allocation_by_unit.append({
            "unit": unit.name,
            "is_core": getattr(unit, "is_core_unit", False),
            "assets_total": unit_assets.count(),
            "assets_assigned": unit_assets.filter(status="assigned").count(),
            "assets_available": unit_assets.filter(status="available").count(),
            "lines_assigned": unit_lines.filter(status="assigned").count(),
        })

    # ── Chart JSON ───────────────────────────────────────────────────────────
    chart_asset_status = json.dumps({
        "labels": ["Available", "Assigned", "Maintenance", "Retired"],
        "datasets": [{
            "data": [available_assets, assigned_assets, maintenance_assets, retired_assets],
            "backgroundColor": ["#198754", "#0d6efd", "#ffc107", "#6c757d"],
        }]
    })

    chart_req_status = json.dumps({
        "labels": ["Pending Manager", "Pending ICT", "Assigned", "Received", "Rejected", "Cancelled"],
        "datasets": [{
            "data": [
                req_status_counts.get("pending_manager", 0),
                req_status_counts.get("pending_ict", 0),
                req_status_counts.get("assigned", 0),
                req_status_counts.get("received", 0),
                req_status_counts.get("rejected", 0),
                req_status_counts.get("cancelled", 0),
            ],
            "backgroundColor": ["#fd7e14", "#0dcaf0", "#0d6efd", "#198754", "#dc3545", "#6c757d"],
        }]
    })

    chart_monthly_requests = json.dumps({
        "labels": [m["label"] for m in monthly_requests],
        "datasets": [{
            "label": "Asset Requests",
            "data": [m["count"] for m in monthly_requests],
            "borderColor": "#0d6efd",
            "backgroundColor": "rgba(13,110,253,0.1)",
            "tension": 0.3,
            "fill": True,
        }]
    })

    chart_monthly_creqs = json.dumps({
        "labels": [m["label"] for m in monthly_creqs],
        "datasets": [{
            "label": "Supply Requests",
            "data": [m["count"] for m in monthly_creqs],
            "borderColor": "#198754",
            "backgroundColor": "rgba(25,135,84,0.1)",
            "tension": 0.3,
            "fill": True,
        }]
    })

    cat_labels = [r["category__name"] for r in assets_by_cat[:10]]
    chart_assets_by_cat = json.dumps({
        "labels": cat_labels,
        "datasets": [
            {
                "label": "Available",
                "data": [r["available"] for r in assets_by_cat[:10]],
                "backgroundColor": "#198754",
            },
            {
                "label": "Assigned",
                "data": [r["assigned"] for r in assets_by_cat[:10]],
                "backgroundColor": "#0d6efd",
            },
            {
                "label": "Retired",
                "data": [r["retired"] for r in assets_by_cat[:10]],
                "backgroundColor": "#6c757d",
            },
        ]
    })

    unit_labels = [r["unit__name"] or "Unallocated" for r in assets_by_unit[:10]]
    chart_assets_by_unit = json.dumps({
        "labels": unit_labels,
        "datasets": [
            {
                "label": "Assigned",
                "data": [r["assigned"] for r in assets_by_unit[:10]],
                "backgroundColor": "#0d6efd",
            },
            {
                "label": "Available",
                "data": [r["available"] for r in assets_by_unit[:10]],
                "backgroundColor": "#198754",
            },
        ]
    })

    chart_line_status = json.dumps({
        "labels": ["Available", "Assigned", "Suspended", "Retired"],
        "datasets": [{
            "data": [
                line_status_counts.get("available", 0),
                line_status_counts.get("assigned", 0),
                line_status_counts.get("suspended", 0),
                line_status_counts.get("retired", 0),
            ],
            "backgroundColor": ["#198754", "#0d6efd", "#ffc107", "#6c757d"],
        }]
    })

    return {
        # Meta
        "generated_at": now,
        "agency": agency,

        # Assets
        "total_assets": total_assets,
        "available_assets": available_assets,
        "assigned_assets": assigned_assets,
        "retired_assets": retired_assets,
        "maintenance_assets": maintenance_assets,
        "asset_status_counts": asset_status_counts,
        "assets_by_cat": assets_by_cat,
        "assets_by_unit": assets_by_unit,
        "eol_assets": eol_assets,
        "issued_assets": issued_assets,
        "pending_returns": pending_returns,

        # Requests
        "total_requests": total_requests,
        "req_status_counts": req_status_counts,
        "requests_by_cat": requests_by_cat,
        "monthly_requests": monthly_requests,
        "avg_fulfillment_days": avg_fulfillment_days,

        # Lines
        "total_lines": total_lines,
        "line_status_counts": line_status_counts,
        "lines_by_type": lines_by_type,
        "line_by_unit": line_by_unit,
        "all_lines": all_lines,
        "reactivation_requests": reactivation_requests,

        # Consumables
        "all_consumables": all_consumables,
        "stock_summary": stock_summary,
        "total_creqs": total_creqs,
        "creq_status_counts": creq_status_counts,
        "top_consumables": top_consumables,
        "monthly_creqs": monthly_creqs,
        "recent_stock_logs": recent_stock_logs,

        # Allocation
        "allocation_by_unit": allocation_by_unit,

        # Charts
        "chart_asset_status": chart_asset_status,
        "chart_req_status": chart_req_status,
        "chart_monthly_requests": chart_monthly_requests,
        "chart_monthly_creqs": chart_monthly_creqs,
        "chart_assets_by_cat": chart_assets_by_cat,
        "chart_assets_by_unit": chart_assets_by_unit,
        "chart_line_status": chart_line_status,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Main report view
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def asset_reports(request):
    user = request.user
    agency = getattr(user, "agency", None)

    if not agency:
        messages.error(request, "You are not assigned to an agency.")
        return redirect("accounts:profile")

    svc, _ = AgencyServiceConfig.objects.get_or_create(agency=agency)
    if not svc.asset_mgmt_enabled and not user.is_superuser:
        messages.warning(request, "Asset Management is not enabled for your agency.")
        return redirect("accounts:profile")

    # Ensure AgencyAssetRoles exists so _require_manager can query it
    AgencyAssetRoles.objects.get_or_create(agency=agency)

    if not _require_manager(user, agency):
        messages.error(
            request,
            "Access restricted. Reports are available to ICT custodians, "
            "Operations Managers, unit heads, and administrators."
        )
        return redirect("accounts:asset_management")

    data = _build_report_data(agency)
    return render(request, "accounts/assets/asset_reports.html", data)


# ─────────────────────────────────────────────────────────────────────────────
# Excel download
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def asset_reports_excel(request):
    """Download a comprehensive multi-sheet Excel workbook."""
    user = request.user
    agency = getattr(user, "agency", None)

    if not agency:
        messages.error(request, "Access denied.")
        return redirect("accounts:asset_management")

    AgencyAssetRoles.objects.get_or_create(agency=agency)

    if not _require_manager(user, agency):
        messages.error(request, "Access denied.")
        return redirect("accounts:asset_management")

    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side, numbers
        )
        from openpyxl.utils import get_column_letter
        from openpyxl.chart import BarChart, Reference, PieChart
        from openpyxl.chart.series import DataPoint
    except ImportError:
        messages.error(request, "openpyxl is required for Excel export. Run: pip install openpyxl")
        return redirect("accounts:asset_management")

    data = _build_report_data(agency)
    wb = openpyxl.Workbook()

    # ── Styles ────────────────────────────────────────────────────────────────
    BLUE = "1F3864"
    LIGHT_BLUE = "D6E4F0"
    GREEN = "1E7E34"
    LIGHT_GREEN = "D4EDDA"
    AMBER = "856404"
    LIGHT_AMBER = "FFF3CD"
    RED = "721C24"
    LIGHT_RED = "F8D7DA"
    WHITE = "FFFFFF"
    GREY = "F8F9FA"

    def header_font(color=WHITE, bold=True, size=11):
        return Font(name="Calibri", bold=bold, color=color, size=size)

    def header_fill(color=BLUE):
        return PatternFill("solid", fgColor=color)

    def body_font(bold=False, color="000000", size=10):
        return Font(name="Calibri", bold=bold, color=color, size=size)

    def body_fill(color=WHITE):
        return PatternFill("solid", fgColor=color)

    def thin_border():
        s = Side(style="thin", color="CCCCCC")
        return Border(left=s, right=s, top=s, bottom=s)

    def center():
        return Alignment(horizontal="center", vertical="center", wrap_text=True)

    def left():
        return Alignment(horizontal="left", vertical="center", wrap_text=True)

    def _write_headers(ws, headers, row=1, fill_color=BLUE, font_color=WHITE):
        for col_idx, (title, width) in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=title)
            cell.font = header_font(color=font_color)
            cell.fill = header_fill(color=fill_color)
            cell.alignment = center()
            cell.border = thin_border()
            ws.column_dimensions[get_column_letter(col_idx)].width = width
        ws.row_dimensions[row].height = 22

    def _write_row(ws, row_idx, values, alt=False):
        fill = body_fill(GREY if alt else WHITE)
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = body_font()
            cell.fill = fill
            cell.alignment = left()
            cell.border = thin_border()
        ws.row_dimensions[row_idx].height = 18

    def _title_row(ws, title, col_span, row=1, color=BLUE):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_span)
        cell = ws.cell(row=row, column=1, value=title)
        cell.font = Font(name="Calibri", bold=True, color=WHITE, size=14)
        cell.fill = PatternFill("solid", fgColor=color)
        cell.alignment = center()
        ws.row_dimensions[row].height = 30

    def _sub_title(ws, text, col_span, row=2, color=LIGHT_BLUE):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_span)
        cell = ws.cell(row=row, column=1, value=text)
        cell.font = Font(name="Calibri", bold=False, color="333333", size=10, italic=True)
        cell.fill = PatternFill("solid", fgColor=color)
        cell.alignment = left()
        ws.row_dimensions[row].height = 16

    generated = data["generated_at"].strftime("%d %b %Y %H:%M")
    agency_name = str(agency)

    # ── Sheet 1: Summary ─────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_view.showGridLines = False
    _title_row(ws, f"{agency_name} — Asset Management Report", 4)
    _sub_title(ws, f"Generated: {generated}", 4)

    summary_data = [
        ("ASSETS", "", "", ""),
        ("Total Assets", data["total_assets"], "", ""),
        ("Available", data["available_assets"], "", ""),
        ("Assigned / In Use", data["assigned_assets"], "", ""),
        ("Under Maintenance", data["maintenance_assets"], "", ""),
        ("Retired / Disposed", data["retired_assets"], "", ""),
        ("EOL / Due for Replacement", len(data["eol_assets"]), "", ""),
        ("", "", "", ""),
        ("ASSET REQUESTS", "", "", ""),
        ("Total Requests", data["total_requests"], "", ""),
        ("Pending Manager Approval", data["req_status_counts"].get("pending_manager", 0), "", ""),
        ("Pending ICT Assignment", data["req_status_counts"].get("pending_ict", 0), "", ""),
        ("Received & Verified", data["req_status_counts"].get("received", 0), "", ""),
        ("Cancelled / Rejected", (data["req_status_counts"].get("cancelled", 0) + data["req_status_counts"].get("rejected", 0)), "", ""),
        ("Avg. Fulfillment Time (days)", data["avg_fulfillment_days"] or "N/A", "", ""),
        ("", "", "", ""),
        ("MOBILE LINES", "", "", ""),
        ("Total Lines", data["total_lines"], "", ""),
        ("Assigned", data["line_status_counts"].get("assigned", 0), "", ""),
        ("Available", data["line_status_counts"].get("available", 0), "", ""),
        ("Suspended", data["line_status_counts"].get("suspended", 0), "", ""),
        ("", "", "", ""),
        ("CONSUMABLES / SUPPLIES", "", "", ""),
        ("Total Supply Requests", data["total_creqs"], "", ""),
        ("Pending Approval", data["creq_status_counts"].get("pending", 0), "", ""),
        ("Fulfilled", data["creq_status_counts"].get("fulfilled", 0), "", ""),
    ]

    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 18

    for r_idx, row_vals in enumerate(summary_data, 3):
        label, val, _, __ = row_vals
        if not label and not val:
            ws.row_dimensions[r_idx].height = 8
            continue
        if val == "":  # section header
            cell = ws.cell(row=r_idx, column=1, value=label)
            cell.font = Font(name="Calibri", bold=True, color=WHITE, size=11)
            cell.fill = PatternFill("solid", fgColor="2E4057")
            ws.merge_cells(start_row=r_idx, start_column=1, end_row=r_idx, end_column=4)
            cell.alignment = left()
            ws.row_dimensions[r_idx].height = 20
        else:
            c1 = ws.cell(row=r_idx, column=1, value=label)
            c2 = ws.cell(row=r_idx, column=2, value=val)
            c1.font = body_font()
            c2.font = body_font(bold=True)
            c1.alignment = left()
            c2.alignment = center()
            c1.fill = body_fill(GREY if r_idx % 2 == 0 else WHITE)
            c2.fill = body_fill(GREY if r_idx % 2 == 0 else WHITE)
            c1.border = thin_border()
            c2.border = thin_border()
            ws.row_dimensions[r_idx].height = 18

    ws.freeze_panes = "A3"

    # ── Sheet 2: Asset Inventory ──────────────────────────────────────────────
    ws2 = wb.create_sheet("Asset Inventory")
    ws2.sheet_view.showGridLines = False
    headers2 = [
        ("Asset Name", 28), ("Category", 20), ("Unit", 22),
        ("Status", 14), ("Asset Tag", 16), ("Serial Number", 20),
        ("Current Holder", 24), ("Acquired", 14),
    ]
    _title_row(ws2, "Asset Inventory", len(headers2))
    _sub_title(ws2, f"All registered assets as of {generated}", len(headers2))
    _write_headers(ws2, headers2, row=3)

    assets_qs = Asset.objects.filter(agency=agency).select_related("category", "unit", "current_holder").order_by("category__name", "name")
    for r_idx, a in enumerate(assets_qs, 4):
        holder = a.current_holder.get_full_name() if a.current_holder else "—"
        acquired = a.acquired_at.strftime("%d %b %Y") if getattr(a, "acquired_at", None) else "—"
        _write_row(ws2, r_idx, [
            a.name, a.category.name if a.category else "—",
            a.unit.name if a.unit else "Unallocated",
            a.status.replace("_", " ").title(),
            a.asset_tag or "—", a.serial_number or "—",
            holder, acquired,
        ], alt=(r_idx % 2 == 0))

        # Colour the status cell
        status_cell = ws2.cell(row=r_idx, column=4)
        if a.status == "available":
            status_cell.fill = body_fill(LIGHT_GREEN)
            status_cell.font = body_font(color=GREEN)
        elif a.status == "assigned":
            status_cell.fill = body_fill("D0E8FF")
            status_cell.font = body_font(color="0055AA")
        elif a.status == "retired":
            status_cell.fill = body_fill(LIGHT_RED)
            status_cell.font = body_font(color=RED)
        elif a.status == "maintenance":
            status_cell.fill = body_fill(LIGHT_AMBER)
            status_cell.font = body_font(color=AMBER)

    ws2.freeze_panes = "A4"
    ws2.auto_filter.ref = f"A3:{get_column_letter(len(headers2))}3"

    # ── Sheet 3: Asset Requests ───────────────────────────────────────────────
    ws3 = wb.create_sheet("Asset Requests")
    ws3.sheet_view.showGridLines = False
    headers3 = [
        ("Req #", 8), ("Requester", 24), ("Category", 20), ("Unit", 22),
        ("Status", 18), ("Submitted", 14), ("Approved By", 20), ("Assigned Asset", 24),
    ]
    _title_row(ws3, "Asset Requests", len(headers3))
    _sub_title(ws3, f"All asset requests as of {generated}", len(headers3))
    _write_headers(ws3, headers3, row=3)

    req_qs = AssetRequest.objects.filter(agency=agency).select_related(
        "requester", "category", "unit", "assigned_asset", "manager_approved_by"
    ).order_by("-created_at")
    for r_idx, req in enumerate(req_qs, 4):
        _write_row(ws3, r_idx, [
            req.id,
            req.requester.get_full_name() if req.requester else "—",
            req.category.name if req.category else "—",
            req.unit.name if req.unit else "Unallocated",
            req.status.replace("_", " ").title(),
            req.created_at.strftime("%d %b %Y") if req.created_at else "—",
            req.manager_approved_by.get_full_name() if req.manager_approved_by else "—",
            req.assigned_asset.name if req.assigned_asset else "—",
        ], alt=(r_idx % 2 == 0))

    ws3.freeze_panes = "A4"
    ws3.auto_filter.ref = f"A3:{get_column_letter(len(headers3))}3"

    # ── Sheet 4: Assets Issued (Assignments) ──────────────────────────────────
    ws4 = wb.create_sheet("Assets Issued")
    ws4.sheet_view.showGridLines = False
    headers4 = [
        ("Asset", 28), ("Category", 20), ("Unit", 22),
        ("Holder", 24), ("Asset Tag", 16), ("Serial Number", 20), ("Assigned Since", 14),
    ]
    _title_row(ws4, "Assets Currently Issued", len(headers4), color="1E4D78")
    _sub_title(ws4, f"All assets with status = Assigned as of {generated}", len(headers4), color="D6E4F0")
    _write_headers(ws4, headers4, row=3, fill_color="1E4D78")

    for r_idx, a in enumerate(data["issued_assets"], 4):
        holder = a.current_holder.get_full_name() if a.current_holder else "—"
        # Try to find assignment date from history
        issued_since = "—"
        last_assign = AssetHistory.objects.filter(
            asset=a, event="assigned"
        ).order_by("-created_at").first()
        if last_assign:
            issued_since = last_assign.created_at.strftime("%d %b %Y")

        _write_row(ws4, r_idx, [
            a.name, a.category.name if a.category else "—",
            a.unit.name if a.unit else "Unallocated",
            holder, a.asset_tag or "—", a.serial_number or "—",
            issued_since,
        ], alt=(r_idx % 2 == 0))

    ws4.freeze_panes = "A4"
    ws4.auto_filter.ref = f"A3:{get_column_letter(len(headers4))}3"

    # ── Sheet 5: Mobile Lines ─────────────────────────────────────────────────
    ws5 = wb.create_sheet("Mobile Lines")
    ws5.sheet_view.showGridLines = False
    headers5 = [
        ("MSISDN", 18), ("Type", 14), ("Provider", 16),
        ("Status", 14), ("Assigned To", 24), ("Unit", 20),
        ("SIM Serial", 18), ("Issued", 14), ("Notes", 30),
    ]
    _title_row(ws5, "Mobile Lines", len(headers5), color="145369")
    _sub_title(ws5, f"All mobile lines as of {generated}", len(headers5), color="D0EAF4")
    _write_headers(ws5, headers5, row=3, fill_color="145369")

    for r_idx, line in enumerate(data["all_lines"], 4):
        holder = line.assigned_to.get_full_name() if line.assigned_to else "—"
        unit_name = getattr(getattr(line.assigned_to, "unit", None), "name", "—") if line.assigned_to else "—"
        issued = line.issued_at.strftime("%d %b %Y") if getattr(line, "issued_at", None) else "—"
        _write_row(ws5, r_idx, [
            line.msisdn, line.get_line_type_display() if hasattr(line, "get_line_type_display") else line.line_type,
            line.provider or "—", line.status.replace("_", " ").title(),
            holder, unit_name, line.sim_serial or "—", issued,
            (line.notes or "")[:80],
        ], alt=(r_idx % 2 == 0))

        status_cell = ws5.cell(row=r_idx, column=4)
        if line.status == "assigned":
            status_cell.fill = body_fill("D0E8FF"); status_cell.font = body_font(color="0055AA")
        elif line.status == "available":
            status_cell.fill = body_fill(LIGHT_GREEN); status_cell.font = body_font(color=GREEN)
        elif line.status == "suspended":
            status_cell.fill = body_fill(LIGHT_AMBER); status_cell.font = body_font(color=AMBER)

    ws5.freeze_panes = "A4"
    ws5.auto_filter.ref = f"A3:{get_column_letter(len(headers5))}3"

    # ── Sheet 6: Consumables / Supplies ──────────────────────────────────────
    ws6 = wb.create_sheet("Supplies & Stock")
    ws6.sheet_view.showGridLines = False
    headers6 = [
        ("Item", 28), ("Category", 20), ("Stock Qty", 12),
        ("Low-Stock Threshold", 20), ("Unit of Measure", 16),
        ("Max / Request", 14), ("Status", 16),
    ]
    _title_row(ws6, "Consumables & Supplies Stock", len(headers6), color="145A32")
    _sub_title(ws6, f"Current stock levels as of {generated}", len(headers6), color="D4EDDA")
    _write_headers(ws6, headers6, row=3, fill_color="145A32")

    for r_idx, item in enumerate(data["all_consumables"], 4):
        status = "OK"
        fill_color = WHITE
        font_color = "000000"
        if getattr(item, "is_out_of_stock", False):
            status = "OUT OF STOCK"; fill_color = LIGHT_RED; font_color = RED
        elif getattr(item, "is_low_stock", False):
            status = "LOW STOCK"; fill_color = LIGHT_AMBER; font_color = AMBER

        _write_row(ws6, r_idx, [
            item.name, item.category.name if item.category else "—",
            item.stock_qty, item.low_stock_threshold,
            getattr(item, "unit_of_measure", "—"),
            getattr(item, "max_per_request", "—"),
            status,
        ], alt=(r_idx % 2 == 0))

        for col in range(1, len(headers6) + 1):
            c = ws6.cell(row=r_idx, column=col)
            c.fill = body_fill(fill_color)
            if col == len(headers6):
                c.font = body_font(bold=(status != "OK"), color=font_color)

    ws6.freeze_panes = "A4"
    ws6.auto_filter.ref = f"A3:{get_column_letter(len(headers6))}3"

    # ── Sheet 7: Supply Requests ──────────────────────────────────────────────
    ws7 = wb.create_sheet("Supply Requests")
    ws7.sheet_view.showGridLines = False
    headers7 = [
        ("Req #", 8), ("Requester", 24), ("Unit", 20), ("Status", 16),
        ("Items (summary)", 40), ("Submitted", 14), ("Approved By", 22),
    ]
    _title_row(ws7, "Consumable / Supply Requests", len(headers7), color="145A32")
    _sub_title(ws7, f"All supply requests as of {generated}", len(headers7), color="D4EDDA")
    _write_headers(ws7, headers7, row=3, fill_color="145A32")

    creqs_qs = ConsumableRequest.objects.filter(agency=agency).select_related(
        "requester", "unit", "approved_by"
    ).prefetch_related("line_items__item").order_by("-created_at")
    for r_idx, creq in enumerate(creqs_qs, 4):
        items_summary = ", ".join(
            f"{ci.item.name} x{ci.quantity_requested}" for ci in creq.line_items.all()[:5]
        )
        _write_row(ws7, r_idx, [
            creq.id,
            creq.requester.get_full_name() if creq.requester else "—",
            creq.unit.name if creq.unit else "—",
            creq.status.replace("_", " ").title(),
            items_summary or "—",
            creq.created_at.strftime("%d %b %Y") if creq.created_at else "—",
            creq.approved_by.get_full_name() if getattr(creq, "approved_by", None) else "—",
        ], alt=(r_idx % 2 == 0))

    ws7.freeze_panes = "A4"
    ws7.auto_filter.ref = f"A3:{get_column_letter(len(headers7))}3"

    # ── Sheet 8: Allocation by Unit ───────────────────────────────────────────
    ws8 = wb.create_sheet("Allocation by Unit")
    ws8.sheet_view.showGridLines = False
    headers8 = [
        ("Unit", 30), ("Core Unit", 12),
        ("Total Assets", 14), ("Assets Assigned", 16), ("Assets Available", 16),
        ("Lines Assigned", 14),
    ]
    _title_row(ws8, "Allocation Summary by Unit", len(headers8), color="2E4057")
    _sub_title(ws8, f"Assets and lines per organisational unit as of {generated}", len(headers8), color="DDE8F5")
    _write_headers(ws8, headers8, row=3, fill_color="2E4057")

    for r_idx, row in enumerate(data["allocation_by_unit"], 4):
        _write_row(ws8, r_idx, [
            row["unit"],
            "Yes" if row["is_core"] else "No",
            row["assets_total"], row["assets_assigned"],
            row["assets_available"], row["lines_assigned"],
        ], alt=(r_idx % 2 == 0))

    ws8.freeze_panes = "A4"

    # ── Finalize ──────────────────────────────────────────────────────────────
    wb.active = ws  # back to Summary

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f"asset_report_{date.today().strftime('%Y%m%d')}.xlsx"
    response = HttpResponse(
        output.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ─────────────────────────────────────────────────────────────────────────────
# Word download
# ─────────────────────────────────────────────────────────────────────────────

@login_required
def asset_reports_word(request):
    """Download a comprehensive Word (.docx) report."""
    user = request.user
    agency = getattr(user, "agency", None)

    if not agency:
        messages.error(request, "Access denied.")
        return redirect("accounts:asset_management")

    AgencyAssetRoles.objects.get_or_create(agency=agency)

    if not _require_manager(user, agency):
        messages.error(request, "Access denied.")
        return redirect("accounts:asset_management")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree
    except ImportError:
        messages.error(request, "python-docx is required for Word export. Run: pip install python-docx")
        return redirect("accounts:asset_management")

    data = _build_report_data(agency)
    doc = DocxDocument()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Style helpers ─────────────────────────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_cell_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for border_name in ("top", "left", "bottom", "right"):
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn("w:val"), "single")
            border_el.set(qn("w:sz"), "4")
            border_el.set(qn("w:space"), "0")
            border_el.set(qn("w:color"), "CCCCCC")
            tcPr.append(border_el)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        if level == 1:
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        elif level == 2:
            p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        return p

    def add_para(text, bold=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_kpi_row(label, value, color_hex="1F3864"):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.font.size = Pt(10)
        r2 = p.add_run(str(value))
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(color_hex)
        return p

    def add_table(headers_widths, rows_data, header_color="1F3864"):
        col_count = len(headers_widths)
        tbl = doc.add_table(rows=1 + len(rows_data), cols=col_count)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = tbl.rows[0]
        for i, (hdr_text, width_in) in enumerate(headers_widths):
            cell = hdr.cells[i]
            cell.width = Inches(width_in)
            set_cell_bg(cell, header_color)
            p = cell.paragraphs[0]
            run = p.add_run(hdr_text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Data rows
        for r_idx, row_vals in enumerate(rows_data):
            row = tbl.rows[r_idx + 1]
            bg = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row_vals):
                cell = row.cells[c_idx]
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(str(val) if val is not None else "—")
                run.font.size = Pt(9)
        return tbl

    generated = data["generated_at"].strftime("%d %B %Y at %H:%M")
    agency_name = str(agency)

    # ── Cover / Title ─────────────────────────────────────────────────────────
    title_p = doc.add_heading(f"{agency_name}", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub_p = doc.add_paragraph("Asset Management — Comprehensive Report")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].bold = True
    sub_p.runs[0].font.size = Pt(14)
    sub_p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    date_p = doc.add_paragraph(f"Generated: {generated}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    add_heading("1. Executive Summary")

    add_heading("1.1 Asset Overview", level=2)
    add_kpi_row("Total Assets Registered", data["total_assets"])
    add_kpi_row("Available", data["available_assets"], "198754")
    add_kpi_row("Assigned / In Use", data["assigned_assets"], "0D6EFD")
    add_kpi_row("Under Maintenance", data["maintenance_assets"], "FFC107")
    add_kpi_row("Retired / Disposed", data["retired_assets"], "6C757D")
    add_kpi_row("EOL / Due for Replacement", len(data["eol_assets"]), "DC3545")

    doc.add_paragraph()
    add_heading("1.2 Asset Requests", level=2)
    add_kpi_row("Total Requests", data["total_requests"])
    add_kpi_row("Pending Manager Approval", data["req_status_counts"].get("pending_manager", 0), "FD7E14")
    add_kpi_row("Pending ICT Assignment", data["req_status_counts"].get("pending_ict", 0), "0DCAF0")
    add_kpi_row("Received & Verified", data["req_status_counts"].get("received", 0), "198754")
    if data["avg_fulfillment_days"]:
        add_kpi_row("Avg. Fulfillment Time", f"{data['avg_fulfillment_days']} days")

    doc.add_paragraph()
    add_heading("1.3 Mobile Lines", level=2)
    add_kpi_row("Total Lines", data["total_lines"])
    add_kpi_row("Assigned", data["line_status_counts"].get("assigned", 0), "0D6EFD")
    add_kpi_row("Available", data["line_status_counts"].get("available", 0), "198754")
    add_kpi_row("Suspended", data["line_status_counts"].get("suspended", 0), "FFC107")

    doc.add_paragraph()
    add_heading("1.4 Consumable Supplies", level=2)
    add_kpi_row("Total Supply Requests", data["total_creqs"])
    add_kpi_row("Pending Approval", data["creq_status_counts"].get("pending", 0), "FD7E14")
    add_kpi_row("Fulfilled", data["creq_status_counts"].get("fulfilled", 0), "198754")

    doc.add_page_break()

    # ── 2. Asset Inventory ────────────────────────────────────────────────────
    add_heading("2. Asset Inventory by Category")
    doc.add_paragraph("Breakdown of registered assets per category, showing availability and assignment status.")
    doc.add_paragraph()

    cat_rows = [
        (r["category__name"] or "Uncategorised", r["total"], r["available"], r["assigned"], r["retired"])
        for r in data["assets_by_cat"]
    ]
    add_table(
        [("Category", 2.2), ("Total", 0.8), ("Available", 1.0), ("Assigned", 1.0), ("Retired", 0.9)],
        cat_rows,
    )

    doc.add_paragraph()
    add_heading("2.1 Asset Inventory by Unit", level=2)
    unit_rows = [
        (r["unit__name"] or "Unallocated", r["total"], r["available"], r["assigned"])
        for r in data["assets_by_unit"]
    ]
    add_table(
        [("Unit", 2.5), ("Total", 0.9), ("Available", 1.0), ("Assigned", 1.0)],
        unit_rows,
    )

    doc.add_paragraph()
    if data["eol_assets"]:
        add_heading("2.2 Assets Due for Replacement (EOL)", level=2)
        eol_rows = [
            (a.name, a.category.name if a.category else "—", a.unit.name if a.unit else "—",
             a.current_holder.get_full_name() if a.current_holder else "—")
            for a in data["eol_assets"]
        ]
        add_table(
            [("Asset", 2.2), ("Category", 1.5), ("Unit", 1.5), ("Holder", 1.7)],
            eol_rows,
            header_color="8B2500",
        )

    doc.add_page_break()

    # ── 3. Assets Issued ──────────────────────────────────────────────────────
    add_heading("3. Assets Currently Issued")
    doc.add_paragraph(f"Total assets assigned to staff: {len(data['issued_assets'])}")
    doc.add_paragraph()

    issued_rows = []
    for a in data["issued_assets"][:50]:  # cap at 50 for readability
        last_assign = AssetHistory.objects.filter(asset=a, event="assigned").order_by("-created_at").first()
        issued_since = last_assign.created_at.strftime("%d %b %Y") if last_assign else "—"
        issued_rows.append((
            a.name, a.category.name if a.category else "—",
            a.unit.name if a.unit else "—",
            a.current_holder.get_full_name() if a.current_holder else "—",
            a.asset_tag or "—", issued_since,
        ))

    add_table(
        [("Asset", 1.8), ("Category", 1.3), ("Unit", 1.3), ("Holder", 1.8), ("Tag", 0.8), ("Since", 0.9)],
        issued_rows,
        header_color="1E4D78",
    )

    if len(data["issued_assets"]) > 50:
        doc.add_paragraph(f"Note: Showing first 50 of {len(data['issued_assets'])} issued assets. Download the Excel report for the full list.")

    doc.add_page_break()

    # ── 4. Mobile Lines ────────────────────────────────────────────────────────
    add_heading("4. Mobile Lines")
    add_heading("4.1 Line Type Breakdown", level=2)
    type_rows = [
        (r["line_type"].replace("_", " + ").title(), r["total"], r["assigned"], r["available"], r["suspended"])
        for r in data["lines_by_type"]
    ]
    add_table(
        [("Type", 1.5), ("Total", 0.8), ("Assigned", 1.0), ("Available", 1.0), ("Suspended", 1.0)],
        type_rows,
        header_color="145369",
    )

    doc.add_paragraph()
    add_heading("4.2 Line Allocation by Unit", level=2)
    lbu_rows = [
        (r["assigned_to__unit__name"] or "Unallocated", r["count"])
        for r in data["line_by_unit"]
    ]
    add_table(
        [("Unit", 3.5), ("Lines Assigned", 1.5)],
        lbu_rows,
        header_color="145369",
    )

    doc.add_page_break()

    # ── 5. Consumables & Supplies ─────────────────────────────────────────────
    add_heading("5. Consumables & Supply Requests")

    add_heading("5.1 Current Stock Levels", level=2)
    consumable_rows = [
        (item.name, item.category.name if item.category else "—",
         item.stock_qty, item.low_stock_threshold,
         getattr(item, "unit_of_measure", "—"),
         "OUT OF STOCK" if getattr(item, "is_out_of_stock", False) else (
             "LOW STOCK" if getattr(item, "is_low_stock", False) else "OK"))
        for item in data["all_consumables"]
    ]
    add_table(
        [("Item", 2.2), ("Category", 1.5), ("Stock", 0.8), ("Threshold", 1.0), ("UoM", 0.9), ("Status", 1.0)],
        consumable_rows,
        header_color="145A32",
    )

    doc.add_paragraph()
    add_heading("5.2 Top Requested Supplies", level=2)
    top_rows = [
        (r["item__name"], r["item__category__name"] or "—", r["total_qty"], r["request_count"])
        for r in data["top_consumables"]
    ]
    add_table(
        [("Item", 2.5), ("Category", 1.5), ("Total Qty Requested", 1.5), ("# Requests", 1.0)],
        top_rows,
        header_color="145A32",
    )

    doc.add_page_break()

    # ── 6. Allocation Summary ─────────────────────────────────────────────────
    add_heading("6. Allocation Summary by Unit")
    alloc_rows = [
        (r["unit"], "Yes" if r["is_core"] else "No",
         r["assets_total"], r["assets_assigned"], r["assets_available"], r["lines_assigned"])
        for r in data["allocation_by_unit"]
    ]
    add_table(
        [("Unit", 2.0), ("Core", 0.6), ("Assets Total", 1.0), ("Assigned", 1.0), ("Available", 1.0), ("Lines", 0.8)],
        alloc_rows,
        header_color="2E4057",
    )

    doc.add_paragraph()
    add_para(f"Report generated by {user.get_full_name() or user.username} on {generated}.", size=9)

    # ── Output ────────────────────────────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"asset_report_{date.today().strftime('%Y%m%d')}.docx"
    response = HttpResponse(
        output.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ─────────────────────────────────────────────────────────────────────────────
# Multi-item create_request handler (patch into view_asset_management.py)
# ─────────────────────────────────────────────────────────────────────────────
# Replace the existing  "if action == 'create_request':"  block with:
#
#   if action == "create_request":
#       from .view_asset_reports import _handle_multi_item_request
#       return _handle_multi_item_request(request, user, agency, svc, roles, _notify_local)
#
# Then keep this function here:

def _handle_multi_item_request(request, user, agency, svc, roles, _notify_local):
    """
    Handle the multi-item asset request form.
    Accepts:
        unit_id       (optional)
        justification (shared across all items)
        items[N][category_id]
        items[N][qty]       (integer, default 1)
        items[N][note]      (per-item note appended to justification)
    Creates one AssetRequest per item row.
    """
    from django.shortcuts import redirect
    from django.contrib import messages

    unit_id = request.POST.get("unit_id")
    justification = (request.POST.get("justification") or "").strip()

    unit = None
    if unit_id:
        try:
            from .models import Unit
            unit = Unit.objects.get(id=unit_id, agency=agency)
        except Unit.DoesNotExist:
            pass
    if not unit:
        unit = getattr(user, "unit", None)

    # Collect item rows from POST
    # Fields arrive as items[0][category_id], items[1][category_id], …
    # We parse them via a loop up to a reasonable max.
    created_ids = []
    errors = []

    idx = 0
    while True:
        cat_id = request.POST.get(f"items[{idx}][category_id]")
        if cat_id is None:
            break
        qty_str = request.POST.get(f"items[{idx}][qty]", "1")
        note = (request.POST.get(f"items[{idx}][note]") or "").strip()

        if not cat_id:
            idx += 1
            continue

        try:
            from .models import AssetCategory, AssetRequest
            from .utils_assets import get_manager_emails_for_request, get_ict_custodian_emails
            category = AssetCategory.objects.get(id=cat_id, agency=agency)
        except AssetCategory.DoesNotExist:
            errors.append(f"Item #{idx + 1}: category not found.")
            idx += 1
            continue

        try:
            qty = max(1, int(qty_str))
        except ValueError:
            qty = 1

        item_justification = justification
        if note:
            item_justification = f"{justification}\n[Item note: {note}]".strip()

        # Create one request per quantity unit? Or one request with qty?
        # We create one request per line (qty stored in justification note).
        # If you have a `qty` field on AssetRequest, pass it; otherwise store in note.
        for _ in range(qty):
            req = AssetRequest.objects.create(
                agency=agency,
                requester=user,
                unit=unit,
                category=category,
                justification=item_justification,
                status="pending_manager" if svc.require_manager_approval else "pending_ict",
            )
            created_ids.append(req.id)

            manager_emails = get_manager_emails_for_request(req)
            _notify_local(
                subject=f"Asset Request #{req.id} — Approval Required",
                to_emails=manager_emails,
                html_template="emails/assets/request_submitted.html",
                ctx={"req": req},
            )
            if req.status == "pending_ict":
                ict_emails = get_ict_custodian_emails(req)
                _notify_local(
                    subject=f"Asset Request #{req.id} — Pending ICT Assignment",
                    to_emails=ict_emails,
                    html_template="emails/assets/request_approved.html",
                    ctx={"req": req, "approved_by": "System (Auto)"},
                )

        idx += 1
        if idx > 50:  # safety limit
            break

    for e in errors:
        messages.error(request, e)

    if created_ids:
        ids_str = ", ".join(f"#{i}" for i in created_ids)
        messages.success(request, f"Asset request(s) submitted: {ids_str}")
    elif not errors:
        messages.warning(request, "No valid items found in your request.")

    from django.shortcuts import redirect
    return redirect("accounts:asset_management")